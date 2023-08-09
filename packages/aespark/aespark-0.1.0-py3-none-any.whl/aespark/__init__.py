from pathlib import Path
from typing import Literal
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
import datetime
import jionlp
import pandas as pd
from docx.shared import Cm
import os
import time
import numpy as np
import jieba
import warnings
import re
import pkg_resources
import unicodedata
from tqdm import tqdm

warnings.filterwarnings('ignore')


class MyDocument(object):
    '''
    功能简介：
        读取模板文档创建空白文档供操作；
    参数解释：
        url :   可选 word模板路径；
    '''
    clolist = ['总流水', '金额']  # 向doc增加表格时，格式化列名在此列表内的列内容，将金额规范为标准格式输出
    def __init__(self, url: str = '', drop: bool = True):
        if url == '':
            dist = pkg_resources.get_distribution("aespark")
            url = f'{dist.location}/aespark/static/word.docx'
        self.doc = Document(url)
        if drop:
            self.all_clean()

    def add_text(self, string: str = '未指定插入内容', level: int = -1):
        '''
        功能简介：
            在docx文档末尾增加新的内容，可以是标题或者段落；
        参数解释：
            string    :     新增的内容；
            level     :     内容格式，默认为段落，其余正整数为标题、对于标题等级；
        '''
        if level == -1:
            self.doc.add_paragraph(string)
        else:
            self.doc.add_heading(string, level=level)

    def add_pic(self, url:str, scale:float=1):
        '''
        功能简介：
            在docx文档末尾插入图片；
        参数解释：
            url 图片文件路径；
            scale   图片缩放比例，默认不缩放；例如 3：放大3倍，0.5：缩小50%；
        '''
        picture = self.doc.add_picture(url)
        picture.width = int(picture.width * scale)
        picture.height = int(picture.height * scale)

    def add_table(self, df: pd.DataFrame, sty: str = 'Grid Table 5 Dark Accent 1', fontsize: int = 10):
        '''
        功能简介：
            在docx文档末尾插入新表格；
        参数解释：
            df          :   需插入的表格；
            sty         :   可选 表格风格；
            fontsize    :   可选 表格内字体大小；
        '''
        def get_fontwidth(text):
            count = 0
            for c in text:
                if unicodedata.east_asian_width(c) in 'FWA':
                    count += 2
                else:
                    count += 0.83
            return count

        def set_tablewidth(table, widths):
            """表格分别设置列宽，单位为Cm"""
            for x, width in enumerate(widths):
                for cell in table.columns[x].cells:
                    cell.width = Cm(width)
        
        for col in self.clolist:
            if col in df.columns.to_list():
                df[col] = df[col].apply(dc_amWrite)

        table = self.doc.add_table(df.shape[0]+1, df.shape[1])
        table.style = sty
        table.style.font.size = Pt(fontsize)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j in range(df.shape[-1]):
            table.cell(0, j).text = df.columns[j]
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                table.cell(i+1, j).text = str(df.values[i, j])
        for col in table.columns:
            for cell in col.cells:
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        widlist = []
        for i in df.columns:
            wid = get_fontwidth(str(i))
            for index in df.index:
                lins = get_fontwidth(str(df[i][index]))
                wid = lins if lins > wid else wid
            widlist.append(wid)
        widlist = [i/sum(widlist)*15.26 for i in widlist]
        set_tablewidth(table, widlist)  # 共15.26

    def add_amLis(self, tt:str):
        '''
        功能介绍：
            添加需要进行金额格式化的列名；
        '''
        self.clolist.append(tt)

    def save(self, url: str = f"./未设置保存名称.docx"):
        '''
        功能简介：
            保存docx文档；
        参数解释：
            url :   str.    保存路径（含文件名及后缀），默认当前程序位置；
        '''
        self.doc.save(url)

    def move_pars(self):
        '''
        功能简介：
            移除所有段落；
        '''
        for par in self.doc.paragraphs:
            par._element.getparent().remove(par._element)

    def move_tables(self):
        '''
        功能简介：
            移除所有表格；
        '''
        for tab in self.doc.tables:
            tab._element.getparent().remove(tab._element)

    def all_clean(self):
        '''
        功能简介：
            清空文档；
        '''
        self.move_pars()
        self.move_tables()

def dc_amWrite(coun):
    '''格式化金额'''
    try:
        coun = format(float('{:.2f}'.format(coun)), ',')
        coun[coun.index('.')+1:]
        coun = coun if len(coun[coun.index('.')+1:]) == 2 else coun+'0'
    finally:
        return coun

def dc_excelAddT(df: pd.DataFrame):
    '''
    功能简介：
        添加'\\t'便于存为csv文件；
    '''
    df.columns = [str(i)+'\t' if str(i).isdigit()
                  and len(str(i)) > 15 else i for i in df.columns]
    for i in df.columns:
        df[i] = df[i].apply(lambda x: str(
            x)+'\t' if str(x).isdigit() and len(str(x)) > 15 else x)

    return df

def dc_csvDelT(df: pd.DataFrame):
    '''
    功能简介：
        删除'\\t'便于后续操作；
    '''
    df.columns = [str(i).replace('\t', '') for i in df.columns]
    for i in df.columns:
        df[i] = df[i].apply(lambda x: str(x).replace('\t', ''))

    return df

def dc_invisCharDel(chars: str | pd.DataFrame, df: bool = False):
    '''
    功能简介：
        清除不可见字符；
    参数解释：
        chars       可传字符可传表，默认传的字符；
        df          如果要传dataframe，该项参数需要填写为 True；
    '''
    if df:
        for i in chars.columns:
            chars[i] = chars[i].apply(lambda x: re.sub(
                u'[\u2000-\u200f\u2028-\u202f\u205f-\u206e]', '', x) if type(x) == str else x)

        return chars
    else:
        if type(chars) == str:
            chars = re.sub(
                u'[\u2000-\u200f\u2028-\u202f\u205f-\u206e]', '', chars)

        return chars

def union_sheet(url: str, limit: int = 42000000, link: bool = False, lex: str = '*'):
    '''
    功能简介：
        合并单个sheet的文件；
    参数解释：
        url     目标文件夹路径(文件夹中只能有.csv.xls或.xlsx格式的文件)；
        limit   输出表容量（多少条数据存一张表；默认80万行）；
        link    是否需要添加数据来源，默认不添加；
        lex     需要合并的文件后缀，默认为所有；
    '''
    files = pd.DataFrame(columns=['文件名称', '文件路径'])
    geshi = pd.DataFrame(columns=['总文件名', '表格式', '文件数量'])

    for i in Path(url).rglob(f'*.{lex}'):
        files.loc[len(files)] = [Path(i).stem, i]

    for i in tqdm(files.index, desc='数据提取'):

        if lex == '*':
            filelex = str(files['文件路径'][i])
            iii = filelex[filelex.rindex('.'):]
            if 'xls' in iii or 'xlsx' in iii:
                df = pd.read_excel(files['文件路径'][i], dtype='str')
            elif 'csv' in iii:
                df = pd.read_csv(files['文件路径'][i],
                                 dtype='str', encoding='gb18030')
        elif lex in ['xls', 'xlsx']:
            df = pd.read_excel(files['文件路径'][i], dtype='str')
        elif lex == 'csv':
            df = pd.read_csv(files['文件路径'][i], dtype='str', encoding='gb18030')

        if link:
            df['原始文件路径'] = files['文件路径'][i]
        lis = df.columns.to_list()
        lis.sort()
        lis = ''.join(lis)

        if lis in list(geshi['表格式']):
            row_index = geshi[geshi['表格式'] == lis].index.tolist()[0]
            geshi['总文件名'][row_index] += files['文件名称'][i]
            geshi['文件数量'][row_index] += 1
            exec(f"hebin{row_index} = pd.concat([hebin{row_index}, df])")
        else:
            exec(f"hebin{len(geshi)} = df.copy()")
            geshi.loc[len(geshi)] = [files['文件名称'][i], lis, 1]

    geshi['总文件名'] = geshi['总文件名'].str.replace(
        '[^\u4e00-\u9fa5]', '', regex=True)

    for i in tqdm(geshi.index, desc='数据产出'):

        if geshi['总文件名'][i] != '':
            result = jieba.tokenize(geshi['总文件名'][i])
            cutresult = pd.DataFrame(columns=['word', 'start'])
            for tk in result:
                cutresult.loc[len(cutresult)] = [tk[0],tk[1]]
            cutresult = cutresult.pivot_table(index='word', values='start', aggfunc=['count', 'sum']).reset_index()
            cutresult.columns = ['word', 'count', 'start']
            cutresult.sort_values(by=['count', 'start'], ascending=[False, True], inplace=True)
            cutresult = cutresult[cutresult['count']==cutresult['count'].max()]
            file_name = ''.join(list(cutresult['word']))
        else:
            file_name = '未知'

        exec(f"hebin{i}.drop_duplicates(inplace=True)")
        exec(f"hebin{i}.reset_index(drop=True, inplace=True)")
        exec(f"hebin{i} = dc_invisCharDel(hebin{i}, df=True)")
        exec(f"hebin{i} = dc_excelAddT(hebin{i})")

        n = ''
        num = geshi['文件数量'][i]
        all = geshi['文件数量'].sum()
        buildFolder('合并数据产出')
        while eval(f"len(hebin{i})") > limit:
            n = 1 if n == '' else n+1
            exec(
                f"hebin{i}.loc[:{limit}].to_csv(r'合并数据产出\{i+1}.{file_name}({num},总{all}){n}.csv', index=False, encoding='gb18030')")
            exec(f"hebin{i} = hebin{i}.loc[{limit+1}:]")
            exec(f"hebin{i}.reset_index(drop=True, inplace=True)")

        exec(
            f"hebin{i}.to_csv(r'合并数据产出\{i+1}.{file_name}({num},总{all}){n}.csv', index=False, encoding='gb18030')")

def union_sheets(url: str, lex: str = '*', link: bool = False, seq: str = '_', ind: int = None, save:bool = True):
    '''
    功能简介：
        按sheet名称合并多个excel文件；
    功能输出：
        一个含有多个dataframe的数组；
    参数解释：
        url     文件夹路径；
        lex     需要合并的文件后缀，默认为所有，如果目标文件夹内有其他类型文件需要设定；
        link    是否需要标记数据来源；
        seq     文件名称分隔符，默认为"_"；
        ind     来源位于文件名称第几个，起始为"0"，缺省时添加文件名称为来源；
    '''
    files = [str(i) for i in Path(url).rglob(f'*.{lex}')]  # 获取目标文件夹下的所有文件路径
    for file in tqdm(files, desc='数据抽取'):
        df = pd.read_excel(file, sheet_name=None,
                           keep_default_na='', dtype='str')
        sheets = list(df.keys())

        if link:  # 若选择了标记数据来源，则给所有sheet内的数据行添加来源内容
            if ind == None:
                lin = file[file.rfind('\\')+1:file.rfind('.')]
            else:
                lin = file[file.rfind('\\')+1:file.rfind('.')].split(seq)[ind]
            for i in range(len(sheets)):
                df[sheets[i]]['数据来源'] = lin

        if 'first' not in locals():
            alldata = df  # 存放所有sheet，用于最终输出
            first = 1
        else:
            for sheet in sheets:  # 按sheet合并
                try:
                    alldata[sheet] = pd.concat([alldata[sheet], df[sheet]])
                except:
                    alldata[sheet] = pd.DataFrame()
                    alldata[sheet] = pd.concat([alldata[sheet], df[sheet]])

    for sheet in list(alldata.keys()):
        alldata[sheet].drop_duplicates(inplace=True)
        alldata[sheet].reset_index(inplace=True, drop=True)

    if save == False:
        print(f"合并完成：共得到{len(alldata.keys())}个表格：{'、'.join([i for i in alldata.keys()])}")
    else:
        buildFolder('多sheet合并数据产出')
        with pd.ExcelWriter(f"多sheet合并数据产出/多sheet合并结果.xlsx") as score_file:
            for sheet in list(alldata.keys()):
                alldata[sheet].to_excel(score_file,sheet_name = sheet, index = False)
    
    return alldata

def parse_idber(idCardNumber: str):
    '''
    功能简介：
        身份证号码解析；
    所需参数：
        idCardNumber    身份证号码；
    输出信息：
        list.  性别、年龄、省、市、区；
    调用示例：
        dataframe[['性别', '年龄', '归属省', '归属市', '归属区']] = dataframe['身份证号'].apply(pd.Series(parse_idber))；
    '''
    try:
        res = jionlp.parse_id_card(idCardNumber)
        province = res['province'] if res['province'] != None else ''
        city = res['city'] if res['city'] != None else ''
        county = res['county'] if res['county'] != None else ''
        age = datetime.datetime.now().year - int(res['birth_year'])
        gender = res['gender']
    except:
        province = city = county = age = gender = ''
    finally:
        return [gender, age, province, city, county]

def parse_phoneber(phoneNumber: str):
    '''
    功能简介：
        手机号码解析；
    所需参数：
        phoneNumber 手机号码；
    return：
        Series.  省、市、运营商；
    调用示例：
        dataframe[['省', '市', '运营商']] = dataframe['手机号'].apply(pd.Series(parse_phoneber))；
    '''
    try:
        res = jionlp.cell_phone_location(phoneNumber)
        province = res['province']
        city = res['city']
        operator = res['operator']
    except:
        province = city = operator = ''
    finally:
        return [province, city, operator]

def buildFolder(url: str):
    '''
    功能简介：
        根据传入路径创建文件夹，自动跳过已存在文件夹；
    所需参数：
        路径，可以是多层文件夹；
    return：
        路径创建情况；
    '''
    if '\\' in url:
        url = url.replace('\\', '/')

    for i in url.split('/'):
        if 'urlstr' not in locals():
            urlstr = i
        else:
            urlstr += '/'+i
        if os.path.exists(urlstr) is False:
            os.mkdir(urlstr)

def dc_amClean(df:pd.DataFrame, clo:str):
    '''
    功能简介：
        交易金额清洗；
    阐述解释：
        df  需要清洗的表；
        clo 需要清洗的列名；
    '''
    try:
        df[clo] = df[clo].astype('float')
    except:
        df.reset_index(drop=True, inplace=True)
        count = 0
        for i in df.index:
            try:
                float(df[clo][i])
            except:
                df[clo][i] = np.nan
                count += 1

        df[clo] = df[clo].astype('float')
        print(f"金额清洗：共清洗错误金额{count}条，占比：{'{:.2%}'.format(count/len(df))}")
    finally:
        return df

def dc_inOutClean(str: str):
    '''
    功能简介：
        统一借贷标志；
    所需参数：
        需要清洗的字符，建议配合pandas.apply使用；
    当前可清洗内容：
        出 = ['借', '出', '支出', '付', 'D']；
        进 = ['贷', '进', '收入', '收', 'C']；
    如果发现了新的借贷标志可以进行添加；
    '''
    jie = ['借', '出', '支出', '付', 'D']
    dai = ['贷', '进', '收入', '收', 'C']
    if str in jie:
        return '出'
    if str in dai:
        return '进'
    return '其他'

def dc_tryTime(timestr: str, format: str):
    '''
    功能简介：
        格式化时间格式；
    所需参数：
        timestr 需要格式化的字符串；
        format  字符串的格式（%Y年 %m月 %d日 %H时 %M分 %S秒）；
    return：
        清洗成功的时间格式（示例 2023.07.25 16:11:52）；
        若清洗失败则会返回False；
    '''
    try:
        timeStruct = time.strptime(timestr, format)
        times = time.strftime("%Y.%m.%d %H:%M:%S", timeStruct)
        return times
    except:
        return False

def dc_Time(timestr: str):
    '''
    功能简介：
        兼容格式，批量格式化时间格式；
    所需参数：
        timestr 需要格式化的字符串（建议配合pandas.apply使用）；
    return：
        清洗成功的时间格式（示例 2023.07.25 16:11:52）；
        若清洗失败则会返回 nan；
    '''
    if timestr.isdigit():
        if len(timestr) == 14:
            times = dc_tryTime(timestr, '%Y%m%d%H%M%S')
        elif len(timestr) == 12:
            times = dc_tryTime(timestr, '%Y%m%d%H%M')
        elif len(timestr) == 8:
            times = dc_tryTime(timestr, '%Y%m%d')
        else:
            times = dc_tryTime(timestr, '%Y%m%d%H%M%S')
            if times is False:
                times = dc_tryTime(timestr, '%Y%m%d%H%M')
            if times is False:
                times = dc_tryTime(timestr, '%Y%m%d')

    else:
        if '-' in timestr:
            str = '-'
        elif '/' in timestr:
            str = '/'
        elif '.' in timestr:
            str = '.'
        else:
            str = ''
        times = dc_tryTime(timestr, f'%Y{str}%m{str}%d %H:%M:%S')

        if times is False:
            times = dc_tryTime(timestr, f'%Y{str}%m{str}%d %H:%M')

        if times is False:
            times = dc_tryTime(timestr, f'%Y{str}%m{str}%d')

        if times is False and len(timestr) == 26:  # 2016-01-21-21.17.03.704713
            times = dc_tryTime(timestr[:-7], '%Y-%m-%d-%H.%M.%S')

    return times if times else np.nan

class IP():
   
    def __init__(self):
        dist = pkg_resources.get_distribution("aespark")
        self.dfv4 = pd.read_table(f'{dist.location}/aespark/static/ipv4.txt', keep_default_na='')
        self.dfv6 = pd.read_table(f'{dist.location}/aespark/static/ipv6.txt', keep_default_na='')

    def parse_ip(self, ipstr: str):
        '''
        功能简介：
            ipv4及ipv6地址解析；
        return：
            一个字典，包含省、市、区、运营商、地址、原始信息；
        '''

        def fail():
            dic = {
                'statu': 'fail',
                'province': '',
                'city': '',
                'county': '',
                'operators': '',
                'fulladdress': '',
                'fullinformation': '',
            }
            return dic
            
        def parse_ipv4(ipstr):
            IIP = lambda x:sum([256 ** i * int(j)for i, j in enumerate(x.split('.')[::-1])])
            try:
                ind = list(self.dfv4[self.dfv4['起点长整型'] <= IIP(ipstr)].index)[-1]
                dic = {
                    'statu': 'success',
                    'province': self.dfv4['省'][ind],
                    'city': self.dfv4['市'][ind],
                    'county': self.dfv4['区'][ind],
                    'operators': self.dfv4['运营商'][ind],
                    'fulladdress': self.dfv4['地址'][ind],
                    'fullinformation': self.dfv4['原始信息'][ind],
                }
            except:
                dic = fail()

            return dic
        
        def parse_ipv6(ipstr):
            try:
                ind = list(self.dfv6[self.dfv6['ip_start'] <= ipstr].index)[-1]
                dic = {
                    'statu': 'success',
                    'province': self.dfv6['province'][ind],
                    'city': self.dfv6['city'][ind],
                    'county': self.dfv6['area'][ind],
                    'operators': '',
                    'fulladdress': self.dfv6['address'][ind],
                    'fullinformation': self.dfv6['address'][ind]+self.dfv6['location'][ind],
                }
            except:
                dic = fail()
            
            return dic

        ipstr = re.sub('[^0-9a-f:.]', '', ipstr)
        if ':' in ipstr:
            return parse_ipv6(ipstr)
        elif '.' in ipstr:
            return parse_ipv4(ipstr)
        else:
            return parse_ipv4(ipstr)

def piv_transView(df:pd.DataFrame, collist:list):
    '''
    功能简介：
        基于交易数据产出账户交易概况表；
    参数解释：
        df  交易数据表；
        collist 字段列表，顺序需要一致，[主端账户，交易金额，借贷标志，交易时间]；
    调用示例：
        pivdf = piv_transView(df,['用户ID', '金额(元)', '收/支', '创建时间'])；
    '''
    df = df[[collist[0],collist[1], collist[2], collist[3]]]
    df = dc_amClean(df, collist[1])
    df[collist[2]] = df[collist[2]].apply(dc_inOutClean)
    df[collist[3]] = df[collist[3]].apply(dc_Time)
    if '进' not in list(df[collist[2]].drop_duplicates()) or '出' not in list(df[collist[2]].drop_duplicates()):
        df.loc[len(df)] = ['填充用账户', 0, '进', '1999.01.01 11:14:42']
        df.loc[len(df)] = ['填充用账户', 0, '出', '1999.01.01 11:14:42']
    piv = df.pivot_table(index=collist[0], columns=collist[2], values=collist[1], aggfunc=['sum', 'count']).reset_index()
    piv.columns = [''.join([i[1],i[0]]).replace('sum', '金额').replace('count', '次数') for i in piv.columns]
    if '填充用账户' in list(piv[collist[0]]):
        piv.drop(piv[piv[collist[0]]=='填充用账户'].index, inplace=True)
    if len(piv.columns) == 7:
        piv = piv.take([0, 1, 4, 2, 5, 3, 6], axis=1)
    else:
        piv = piv.take([0, 1, 3, 2, 4], axis=1)
    piv.fillna(0, inplace=True)
    piv['总金额'] = piv['其他金额']+piv['出金额']+piv['进金额'] if len(piv.columns) == 7 else piv['出金额']+piv['进金额']
    df[collist[3]] = pd.to_datetime(df[collist[3]])
    piv2 = df.pivot_table(index=collist[0], values=collist[3], aggfunc=['min', 'max']).reset_index()
    piv2.columns = [collist[0], '首次交易', '末次交易']
    piv2['首次交易'] = piv2['首次交易'].apply(lambda x:str(x)[:str(x).index(' ')].replace('-','.'))
    piv2['末次交易'] = piv2['末次交易'].apply(lambda x:str(x)[:str(x).index(' ')].replace('-','.'))

    piv = pd.merge(left=piv, right=piv2, how='left', on=collist[0])
    for i in piv.columns:
        if '金额' in i:
            piv[i] = piv[i].apply(lambda x: '%.2f'%x)
        elif '次数' in i:
            piv[i] = piv[i].astype('int')
    piv.sort_values(by=['总金额', '末次交易', '首次交易'], ascending=[False, False, True], inplace=True)

    return piv

class BANK():

    def __init__(self):
        dist = pkg_resources.get_distribution("aespark")
        self.df = pd.read_table(f'{dist.location}/aespark/static/bankBIN.txt', keep_default_na='')

    def parse_bankber(self, card: str):
        '''
        功能简介：
            银行卡归属行查询；
        return：
            返回一个字典，查询状态、卡类型、归属行；
        '''
        lins = self.df[self.df['长度']==len(card)]
        for i in lins.index:
            if card[:lins['BIN长度'][i]] == str(lins['卡头'][i]):
                return {'statu':'success','lex':lins['卡类型'][i],'bank':lins['银行'][i]}
        return {'statu':'fail','lex':'','bank':''}
    
class Palette():

    def __init__(self) -> None:
        self.df = pd.DataFrame(columns=['主端'])

    def add_color(self, main:str, col:str|list, color:str|int|float|list):
        '''
        功能简介：
            在情况汇总表上作上一个记录
        参数解释：
            main    记到谁的头上
            col 记录的项目名称是什么（即列名），可以是列表，但必须与color数量一致；
            color   需要记录的内容是什么，可以是列表，但必须与col数量一致；
        使用举例：
            add_color('李四', '花呗记录', '有')；
            add_color('张三', ['余额宝', '花呗记录'], [986.32, '有'])；
        '''
        if type(col) != list:
            col = [col]
            color = [color]
        if len(col) != len(color):
            raise Exception('项目(col)与内容(color)的数量需要一致')
        if main not in list(self.df['主端']):
            self.df.loc[len(self.df)] = {'主端':main}
        for i in range(len(col)):
            if col[i] not in self.df.columns.to_list():
                self.df = pd.concat([self.df, pd.DataFrame(columns=[col[i]])])
            ind = list(self.df[self.df['主端']==main].index)[0]
            self.df[col[i]][ind] = color[i]

    def show(self, what:Literal["color", "main", "palette"]='palette'):
        '''
        功能简介：
            展示当前情况汇总表内容
        参数解释：
            what    查看汇总表什么内容；color为记录的项目名称，main为记录主端列表，默认显示整表；
        '''
        if what == 'color':
            print(list(self.df.columns)[1:])
        elif what == 'main':
            print(list(self.df['主端']))
        else:
            print(self.df)

    def save(self, url:str='palette-未设置保存名称.xlsx'):
        self.df['主端'] = self.df['主端'].astype('str')
        self.df.to_excel(url, index=False)