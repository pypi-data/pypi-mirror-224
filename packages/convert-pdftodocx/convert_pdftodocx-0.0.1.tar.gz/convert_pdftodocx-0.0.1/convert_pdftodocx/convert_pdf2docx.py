import os
import re
import fitz
import time
import docx
import calendar
import pythoncom
import comtypes.client


class pdf2docx:

    def __init__(self):
        pass

    def convert_pdf_to_docx(input_file_path, out_path):
        """
        Method Name: convert_pdf_to_docx

        Description: The Given Function Is Use to Convert PDF To Word Document.
                     Also It Will Extract Normal Text & Table Text From Word Document

        Output: Text & Table Text

        Parameters:
            input_file_path :- Path to pdf file to convert
            out_path :- Path to where docx file will be saved

        On Failure: Raise Exception
        Written By: Volody Products
        Version: 1.0
        Revisions: None
        """
        # Docx path
        timeago = calendar.timegm(time.gmtime())
        docpath = os.path.normpath("%s\docfile_%s.docx" % (out_path, timeago))

        try:
            # Attempt using the original code with Word
            pythoncom.CoInitialize()
            word = comtypes.client.CreateObject('Word.Application')
            pythoncom.CoInitialize()
            word.Visible = False
            word.DisplayAlerts = False

            wb = word.Documents.Open(os.path.abspath(input_file_path))
            wb.SaveAs(os.path.abspath(docpath), FileFormat=16)
            wb.Close()
            word.Quit()

            # Rest of the code as before
            doc = docx.Document(docpath)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            table_text = ""
            for table in doc.tables:
                for row in table.rows:
                    row_text = ""
                    for cell in row.cells:
                        row_text += cell.text + "\t"  # Concatenate cell text and add a tab
                    table_text += row_text + "\n"  # Concatenate row text and add a new line

            table_text = re.sub(r'\s+', ' ', table_text)

            return text, table_text

        except Exception:
            pdf_document = fitz.open(input_file_path)
            doc = docx.Document()
            for page_number in range(pdf_document.page_count):
                page = pdf_document.load_page(page_number)
                text = page.get_text()
                doc.add_paragraph(text)

            doc.save(docpath)
            pdf_document.close()

            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            table_text = ""
            for table in doc.tables:
                for row in table.rows:
                    row_text = ""
                    for cell in row.cells:
                        row_text += cell.text + "\t"  # Concatenate cell text and add a tab
                    table_text += row_text + "\n"  # Concatenate row text and add a new line

            table_text = re.sub(r'\s+', ' ', table_text)

            return text, table_text
