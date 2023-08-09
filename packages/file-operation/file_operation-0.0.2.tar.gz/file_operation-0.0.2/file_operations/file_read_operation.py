# -*- coding:utf-8 -*-
"""
@Time : 2023/3/22
@Author : skyoceanchen
@TEL: 18916403796
@File : other_operation.py 
@PRODUCT_NAME : PyCharm 
"""
# from win32com import client as wc  # pip install pywin32
from win32com.client import Dispatch
import docx  # pip install python-docx
import pdfplumber  # pip install pdfplumber


class ReadDOCDocxPdf(object):
    def read_docx(self, path):
        doc = docx.Document(path)  # 绝对路径
        # 读取表格外全部内容
        text_list = []
        for i in doc.paragraphs:
            text = i.text.replace("—", "")
            if text:
                text_list.append(text)
        text = ''.join(text_list)
        return text

    def read_pdf(self, path):
        pdf = pdfplumber.open(path)
        text_list = []
        # 解析pdf全部内容
        for i in pdf.pages:
            text = i.extract_text().replace("\u3000", "").replace("\xa0", "").replace("—", "").replace("(cid:122)", '')
            # str1 += i.extract_text().replace('', '')
            text_list.append(text)
        text = ''.join(text_list)
        return text

    def read_doc(self, path):
        word = Dispatch('Word.Application')  # 打开word应用程序
        word.Visible = 0  # 后台运行,不显示
        word.DisplayAlerts = 0  # 不警告
        doc = word.Documents.Open(FileName=path, Encoding='gbk')
        text_list = []
        for para in doc.paragraphs:
            text = para.Range.Text.replace("\x15", "").replace("\x01", "").replace("\x0c", "").replace("\x07",
                                                                                                       "").replace(
                "\r", "").replace("\x0e", "")
            # text = ''.join(para.Range.Text.strip(" ").split('\r'))
            if text:
                text_list.append(text)
        doc.Close()
        word.Quit()
        text = "".join(text_list)
        return text


class ReadDat():
    def read_dat(self, path):
        with open(path, 'r', encoding='utf-8') as f:
            text = f.read()
        return text
